"""
RCM Redrafting & Design Review Tool — Multi-Provider (Groq + Gemini)
Run with: streamlit run rcm_tool_multi_provider.py

Install dependencies:
    pip install streamlit pandas openpyxl groq pydantic reportlab python-docx google-generativeai
"""

import io
import json
import time
import pandas as pd
import streamlit as st
from pydantic import BaseModel, Field

# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------
st.set_page_config(page_title="RCM AI Tool", layout="wide")
st.title("📊 Risk & Control Matrix (RCM) AI Tool")
st.caption("Powered by Groq / Gemini AI  |  Big 4 Audit Methodology  |  ICFR / SOX / IFC")

# ---------------------------------------------------------------------------
# Sidebar — Provider & API Settings
# ---------------------------------------------------------------------------
st.sidebar.header("⚙️ AI Provider Settings")

provider = st.sidebar.selectbox(
    "AI Provider",
    ["Groq", "Gemini (Google)"],
    help="Select the AI provider. Each requires its own API key."
)

# Auto-read secrets if available
_groq_secret   = st.secrets.get("GROQ_API_KEY",   "") if hasattr(st, "secrets") else ""
_gemini_secret = st.secrets.get("GEMINI_API_KEY",  "") if hasattr(st, "secrets") else ""

if provider == "Groq":
    api_key = st.sidebar.text_input(
        "Groq API Key",
        value=_groq_secret,
        type="password",
        help="Get your free key from https://console.groq.com/keys"
    )
    MODEL_OPTIONS = {
        "Llama 3.3 70B (Best Quality)": "llama-3.3-70b-versatile",
        "Llama 3.1 8B (Fastest)":       "llama-3.1-8b-instant",
        "Mixtral 8x7B (Balanced)":      "mixtral-8x7b-32768",
        "Gemma 2 9B":                   "gemma2-9b-it",
    }
else:  # Gemini
    api_key = st.sidebar.text_input(
        "Gemini API Key",
        value=_gemini_secret,
        type="password",
        help="Get your key from https://aistudio.google.com/app/apikey"
    )
    MODEL_OPTIONS = {
        "Gemini 2.0 Flash (Best Quality)": "gemini-2.0-flash",
        "Gemini 1.5 Flash (Fast)":          "gemini-1.5-flash",
        "Gemini 1.5 Pro (Advanced)":        "gemini-1.5-pro",
        "Gemini 2.0 Flash Lite (Fastest)":  "gemini-2.0-flash-lite",
    }

selected_label = st.sidebar.selectbox("Model", list(MODEL_OPTIONS.keys()))
selected_model = MODEL_OPTIONS[selected_label]

batch_size   = st.sidebar.slider("Rows per batch", 5, 30, 10, step=5,
                                  help="Smaller = safer for rate limits")
retry_delay  = st.sidebar.slider("Retry delay (seconds)", 5, 60, 15, step=5)
max_retries  = st.sidebar.number_input("Max retries per batch", 1, 5, 3)

# ---------------------------------------------------------------------------
# Sidebar — Tool Mode
# ---------------------------------------------------------------------------
st.sidebar.markdown("---")
st.sidebar.header("🛠️ Tool Mode")
tool_mode = st.sidebar.radio(
    "Select what you want to do:",
    [
        "✏️  Redraft & Classify Controls",
        "🔍  Full RCM Design Effectiveness Review",
    ],
)

# ---------------------------------------------------------------------------
# Sidebar — Process Context
# ---------------------------------------------------------------------------
st.sidebar.markdown("---")
st.sidebar.header("📋 Process Context")

PROCESS_OPTIONS = [
    "Purchase to Pay (P2P)", "Order to Cash (O2C)", "Record to Report (R2R)",
    "Hire to Retire (H2R)", "Inventory & Warehouse Management", "Fixed Assets",
    "Treasury & Cash Management", "Tax & Statutory Compliance",
    "Financial Reporting & Close", "Payroll", "Revenue Recognition",
    "Borrowings & Debt Management", "Capital Expenditure (Capex)", "Other (specify below)",
]
INDUSTRY_OPTIONS = [
    "Steel Manufacturing", "FMCG", "Pharmaceutical", "Automobile / Auto Components",
    "IT / Software Services", "NBFC / Financial Services", "Real Estate & Construction",
    "Power & Energy", "Retail", "Cement & Infrastructure", "Textile", "Chemicals",
    "Other (specify below)",
]
ERP_OPTIONS = [
    "SME Assist", "SAP S/4HANA", "SAP ECC", "Oracle Fusion", "Oracle EBS",
    "Microsoft Dynamics 365", "Tally Prime", "Zoho Books", "QuickBooks",
    "Custom / In-house ERP", "No ERP / Manual", "Other (specify below)",
]
REVIEW_PART_OPTIONS = [
    "All Parts (Full Report)",
    "Part 1  – Process Understanding", "Part 2  – Risk Review",
    "Part 3  – Control Review",        "Part 4  – Missing Controls",
    "Part 5  – Unnecessary Controls",  "Part 6  – Key vs Non-Key Classification",
    "Part 7  – Anti-Fraud Review",     "Part 8  – Segregation of Duties",
    "Part 9  – Control Redrafting",    "Part 10 – Financial Statement Assertions",
    "Part 11 – COSO Mapping",          "Part 12 – Control Testing Procedures",
    "Part 13 – Automation Opportunities", "Part 14 – Final Executive Report",
]

selected_process  = st.sidebar.selectbox("Business Process", PROCESS_OPTIONS)
custom_process    = st.sidebar.text_input("Specify Process")  if selected_process  == "Other (specify below)" else ""
selected_industry = st.sidebar.selectbox("Industry", INDUSTRY_OPTIONS)
custom_industry   = st.sidebar.text_input("Specify Industry") if selected_industry == "Other (specify below)" else ""
selected_erp      = st.sidebar.selectbox("ERP / System", ERP_OPTIONS)
custom_erp        = st.sidebar.text_input("Specify ERP")      if selected_erp      == "Other (specify below)" else ""

selected_parts: list[str] = []
if "Design" in tool_mode:
    st.sidebar.markdown("---")
    st.sidebar.header("📑 Review Scope")
    selected_parts = st.sidebar.multiselect(
        "Select report parts to generate:",
        REVIEW_PART_OPTIONS,
        default=["All Parts (Full Report)"],
    )
    if not selected_parts:
        selected_parts = ["All Parts (Full Report)"]

process_label  = custom_process  if selected_process  == "Other (specify below)" else selected_process
industry_label = custom_industry if selected_industry == "Other (specify below)" else selected_industry
erp_label      = custom_erp      if selected_erp      == "Other (specify below)" else selected_erp

# ---------------------------------------------------------------------------
# Pydantic schema
# ---------------------------------------------------------------------------
class ControlItem(BaseModel):
    original_id: int
    redrafted_control: str = Field(description="Audit-ready control description: Who, What, When, How.")
    anti_fraud: str        = Field(description="'Y' if control prevents/detects fraud, else 'N'.")
    risk_rating: str       = Field(description="'High', 'Medium', or 'Low'.")
    control_type: str      = Field(description="'Key' or 'Non-Key'.")

class RCMBatchResponse(BaseModel):
    items: list[ControlItem]

# ---------------------------------------------------------------------------
# Prompt constants
# ---------------------------------------------------------------------------
BIG4_REDRAFTING_PROMPT = """
=== BIG 4 AUDIT METHODOLOGY — CONTROL REDRAFTING RULES ===
You are an Internal Financial Controls (IFC)/ICFR/RCM specialist.
Rewrite each control description using Big 4 audit methodology.
The redrafted control MUST clearly answer ALL of the following:
• WHO performs the control? • WHAT is performed? • HOW is it performed?
• WHAT is reviewed? • WHO reviews/approves it? • WHEN / Frequency?
• WHAT evidence is generated? • WHAT risk does the control mitigate?
Use this structure: "The <Control Owner> prepares/verifies/reconciles/reviews <document>.
The <Reviewer/Approver> reviews and approves the same after verifying <specific checks>.
Upon approval, the transaction is processed/recorded/filed in <ERP/System>.
This ensures <control objective>."
Rules: formal ICFR language, active voice, no assumptions, no omissions.
=== END OF REDRAFTING RULES ===
"""

ANTI_FRAUD_CRITERIA = """
=== ICFR/SOX ANTI-FRAUD CONTROL CLASSIFICATION CRITERIA ===
Anti-Fraud (Y): prevents/detects unauthorized transactions, management override,
fraudulent financial reporting, misappropriation, fake vendors/customers/employees,
unauthorized payments, master data changes, duplicate/fictitious payments,
journal entry manipulation.
NOT Anti-Fraud (N): ONLY ensures mathematical accuracy, statutory compliance,
routine reconciliations, report completeness, or operational efficiency.
Assess each control independently — do NOT default all to Y or N.
=== END OF ANTI-FRAUD CRITERIA ===
"""

ICFR_SOX_CRITERIA = """
=== ICFR / SOX KEY CONTROL CLASSIFICATION CRITERIA ===
KEY CONTROL if ANY: prevents/detects material misstatement, addresses significant
financial reporting risk, involves management review/approval, relates to statutory
compliance (Income Tax, GST, TDS, Companies Act), no compensating control exists,
covers significant account balance, fraud prevention purpose.
NON-KEY if ALL: operational/administrative only, data prep with no financial reporting
impact, supporting activity with another review control existing, failure won't cause
material financial reporting error.
NEVER classify management review of tax, financial statements, journal entries,
provisions, reconciliations, bank payments, statutory returns as Non-Key unless a
higher-level compensating control exists.
=== END OF CLASSIFICATION CRITERIA ===
"""


def build_redraft_prompt(controls_input, has_risk_col, process, industry, erp):
    risk_instruction = (
        "The 'existing_risk' field contains the pre-defined risk description. "
        "Use it to improve Risk Rating, Anti-Fraud assessment, and Key/Non-Key classification.\n"
        if has_risk_col else ""
    )
    context_block = (
        f"\nENGAGEMENT CONTEXT:\n"
        f"• Business Process : {process or 'Not specified'}\n"
        f"• Industry         : {industry or 'Not specified'}\n"
        f"• ERP / System     : {erp or 'Not specified'}\n"
        if any([process, industry, erp]) else ""
    )
    return f"""You are a senior internal audit expert specialising in ICFR/SOX internal controls.
{context_block}
{risk_instruction}
{BIG4_REDRAFTING_PROMPT}
{ANTI_FRAUD_CRITERIA}
{ICFR_SOX_CRITERIA}

For each control return:
1. REDRAFTED CONTROL: Big 4 style, answering WHO/WHAT/HOW/REVIEW/APPROVAL/SYSTEM/EVIDENCE/OBJECTIVE.
   Reference ERP ({erp or 'the ERP system'}) where applicable.
2. ANTI-FRAUD (Y or N): Apply criteria strictly. Assess independently.
3. RISK RATING (High/Medium/Low): High=material financial impact; Medium=moderate; Low=minor.
4. CONTROL TYPE (Key or Non-Key): Apply ICFR/SOX criteria strictly.

Preserve original_id exactly as given.

Return ONLY a valid JSON object in this exact format (no markdown, no explanation):
{{
  "items": [
    {{
      "original_id": <int>,
      "redrafted_control": "<string>",
      "anti_fraud": "<Y or N>",
      "risk_rating": "<High or Medium or Low>",
      "control_type": "<Key or Non-Key>"
    }}
  ]
}}

Controls to process:
{json.dumps(controls_input, indent=2)}"""


def build_design_review_prompt(process, industry, erp, parts, rcm_text):
    all_parts = "All Parts (Full Report)" in parts
    def include(p): return all_parts or any(p in x for x in parts)

    scope_note = (
        "Generate ALL 14 parts of the Design Effectiveness Review."
        if all_parts else "Generate ONLY the following parts: " + ", ".join(parts) + "."
    )
    parts_block = ""
    if include("Part 1"):
        parts_block += "\n──────────────────────────────\nPART 1 – PROCESS UNDERSTANDING\n──────────────────────────────\nIdentify: Business Process, Process Objective, Major Financial Statement Accounts impacted, Significant Risks, Process Boundaries, Upstream and Downstream Processes.\n"
    if include("Part 2"):
        parts_block += "\n──────────────────────────────\nPART 2 – RISK REVIEW\n──────────────────────────────\nIdentify: Missing Risks, Duplicate Risks, Incorrectly Defined Risks, Risks without Controls, Over-controlled Risks, Risks not linked to Financial Reporting. Give reasons for every observation.\n"
    if include("Part 3"):
        parts_block += "\n──────────────────────────────\nPART 3 – CONTROL REVIEW\n──────────────────────────────\nFor every control evaluate: Does it mitigate the risk? Preventive or detective? Manual/automated/IT-dependent? Complete? Professional RCM language? Measurable and testable? Missing review step? Evidence generated? Appropriate control owner/reviewer/approver? Sufficient approval hierarchy?\n"
    if include("Part 4"):
        parts_block += "\n──────────────────────────────\nPART 4 – IDENTIFY MISSING CONTROLS\n──────────────────────────────\nCompare against industry best practices. Identify missing controls for: Approvals, Authorizations, Reconciliations, Maker-Checker, ERP validations, System controls, Management Reviews, Statutory Compliance, Exception Reports, Master Data Controls, Journal Entries, Bank Controls, Tax Controls.\nFor every missing control provide a table: Risk | Suggested Control | Owner | Reviewer | Frequency | Preventive/Detective | Manual/Automated | Evidence | Key/Non-Key | Anti-Fraud | Assertion | COSO Component.\n"
    if include("Part 5"):
        parts_block += "\n──────────────────────────────\nPART 5 – UNNECESSARY CONTROLS\n──────────────────────────────\nIdentify controls that are operational only, documentation only, duplicate, or administrative with no financial reporting risk reduction. Explain why each is unnecessary.\n"
    if include("Part 6"):
        parts_block += "\n──────────────────────────────\nPART 6 – KEY vs NON-KEY CLASSIFICATION\n──────────────────────────────\nClassify every control as Key or Non-Key using ICFR/SOX criteria. Provide reason for every classification.\n"
    if include("Part 7"):
        parts_block += "\n──────────────────────────────\nPART 7 – ANTI-FRAUD REVIEW\n──────────────────────────────\nIdentify: Existing Anti-Fraud Controls, Missing Anti-Fraud Controls, Fraud Risks, Management Override Risks, SoD Issues, Fake Vendor/Customer Risk, Unauthorized/Duplicate Payments, Manual Journal Risk, Revenue Manipulation.\n"
    if include("Part 8"):
        parts_block += "\n──────────────────────────────\nPART 8 – SEGREGATION OF DUTIES\n──────────────────────────────\nCheck for SoD conflicts: Preparation & Approval, Vendor Creation & Payment, Customer Creation & Credit Approval, Bank Beneficiary Addition & Payment Release, Journal Entry Preparation & Approval, Inventory/Asset Custody & Accounting, Payroll Preparation & Approval. Highlight all conflicts with recommended remediation.\n"
    if include("Part 9"):
        parts_block += "\n──────────────────────────────\nPART 9 – CONTROL REDRAFTING\n──────────────────────────────\nFor every weak or incomplete control, rewrite in Big 4 style specifying WHO | WHAT | HOW | REVIEW | APPROVAL | SYSTEM | EVIDENCE | CONTROL OBJECTIVE. Use formal ICFR language, active voice.\n"
    if include("Part 10"):
        parts_block += "\n──────────────────────────────\nPART 10 – FINANCIAL STATEMENT ASSERTIONS\n──────────────────────────────\nFor every control identify assertions covered: Existence, Completeness, Accuracy, Valuation, Rights & Obligations, Presentation & Disclosure.\n"
    if include("Part 11"):
        parts_block += "\n──────────────────────────────\nPART 11 – COSO MAPPING\n──────────────────────────────\nMap every control to COSO 2013: Control Environment, Risk Assessment, Control Activities, Information & Communication, Monitoring Activities.\n"
    if include("Part 12"):
        parts_block += "\n──────────────────────────────\nPART 12 – CONTROL TESTING PROCEDURES\n──────────────────────────────\nFor every Key Control provide: Testing Procedure (step-by-step), Audit Evidence required, Sample Size, Expected Result, Possible Exceptions.\n"
    if include("Part 13"):
        parts_block += f"\n──────────────────────────────\nPART 13 – AUTOMATION OPPORTUNITIES\n──────────────────────────────\nIdentify controls that can be automated using {erp}/ERP. For each, suggest specific system configuration, workflow rule, or automated control.\n"
    if include("Part 14"):
        parts_block += "\n──────────────────────────────\nPART 14 – FINAL EXECUTIVE REPORT\n──────────────────────────────\nPrepare professional executive report: Executive Summary, Missing/Weak/Duplicate/Unnecessary Controls Summary, SoD Issues, Anti-Fraud Gaps, Key Control Summary, Process Maturity Score (1–5), Overall Design Effectiveness Rating (Effective/Partially Effective/Ineffective), Top 10 Recommendations (ranked), Priority Matrix (Critical/High/Medium/Low). Present all in professional tables.\n"

    return f"""You are a Senior Manager in Risk Advisory (ICFR/IFC/SOX) with 20+ years at Big 4 firms.

ENGAGEMENT CONTEXT:
• Business Process : {process}
• Industry         : {industry}
• ERP / System     : {erp}

BENCHMARK FRAMEWORKS: COSO 2013, IFC (Companies Act 2013), ICAI Guidance Note, SOX, Anti-Fraud Framework.

══════════════════════════════════════════════════════════════════════
CRITICAL READING INSTRUCTIONS — READ BEFORE REVIEWING:
══════════════════════════════════════════════════════════════════════
The RCM data below uses clearly labelled fields for each row:
  • "Risk Description"  = the risk being addressed
  • "Control Activity"  = the EXISTING control already in place at the organisation

RULES YOU MUST FOLLOW:
1. NEVER report a "Control Activity" as a missing control. It already exists.
2. When identifying MISSING controls (Part 4), only flag risks that have
   NO "Control Activity" entry, OR risks that are entirely absent from the RCM
   based on industry best practices for {process} in {industry}.
3. When reviewing controls (Part 3), evaluate the quality of each
   "Control Activity" — do not pretend it doesn't exist.
4. Base ALL observations strictly on the RCM data provided.
   Do not hallucinate controls or risks that are not in the data.
══════════════════════════════════════════════════════════════════════

TASK: {scope_note}
Perform a comprehensive Design Effectiveness Review of the RCM below.
Present all findings in professional tables with clear headings, suitable for an Internal Audit / ICFR report.
Maintain Big 4 consulting style throughout.

{parts_block}

──────────────────────────────
RCM DATA FOR REVIEW:
──────────────────────────────
{rcm_text}
"""


# ---------------------------------------------------------------------------
# Provider abstraction layer
# ---------------------------------------------------------------------------
SYSTEM_PROMPT_REDRAFT = (
    "You are a senior internal audit expert specialising in ICFR/SOX internal controls. "
    "Apply Big 4 audit methodology strictly. Always respond with valid JSON only."
)
SYSTEM_PROMPT_REVIEW = (
    "You are a Senior Manager in Risk Advisory (ICFR/IFC/SOX) with 20+ years at Big 4 firms. "
    "Produce comprehensive, professional RCM Design Effectiveness Review reports. "
    "Always use structured tables, clear headings, and formal audit language."
)


def _clean_json(raw: str) -> str:
    raw = raw.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith("json"):
            raw = raw[4:]
    return raw.strip()


# ── Groq ────────────────────────────────────────────────────────────────────
def _groq_redraft(client, model, prompt, max_retries, retry_delay) -> RCMBatchResponse:
    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT_REDRAFT},
                    {"role": "user",   "content": prompt},
                ],
                temperature=0.2,
                response_format={"type": "json_object"},
            )
            raw = resp.choices[0].message.content.strip()
            return RCMBatchResponse.model_validate(json.loads(_clean_json(raw)))
        except Exception as e:
            last_err = e
            if "429" in str(e) or "rate_limit" in str(e).lower():
                wait = retry_delay * attempt
                st.warning(f"⏳ Groq rate limit (attempt {attempt}/{max_retries}). Waiting {wait}s…")
                time.sleep(wait)
            else:
                raise
    raise Exception(f"Groq: all {max_retries} retries exhausted. Last error: {last_err}")


def _groq_review(client, model, prompt, max_retries, retry_delay) -> str:
    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT_REVIEW},
                    {"role": "user",   "content": prompt},
                ],
                temperature=0.3,
                max_tokens=8000,
            )
            return resp.choices[0].message.content.strip()
        except Exception as e:
            last_err = e
            if "429" in str(e) or "rate_limit" in str(e).lower():
                wait = retry_delay * attempt
                st.warning(f"⏳ Groq rate limit (attempt {attempt}/{max_retries}). Waiting {wait}s…")
                time.sleep(wait)
            else:
                raise
    raise Exception(f"Groq: all {max_retries} retries exhausted. Last error: {last_err}")


# ── Gemini ──────────────────────────────────────────────────────────────────
def _gemini_redraft(client, model_name, prompt, max_retries, retry_delay) -> RCMBatchResponse:
    import google.generativeai as genai
    model = genai.GenerativeModel(
        model_name=model_name,
        system_instruction=SYSTEM_PROMPT_REDRAFT,
        generation_config=genai.GenerationConfig(
            temperature=0.2,
            response_mime_type="application/json",
        ),
    )
    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            resp = model.generate_content(prompt)
            raw  = resp.text.strip()
            return RCMBatchResponse.model_validate(json.loads(_clean_json(raw)))
        except Exception as e:
            last_err = e
            err = str(e).lower()
            if "429" in str(e) or "quota" in err or "rate" in err:
                wait = retry_delay * attempt
                st.warning(f"⏳ Gemini rate limit (attempt {attempt}/{max_retries}). Waiting {wait}s…")
                time.sleep(wait)
            else:
                raise
    raise Exception(f"Gemini: all {max_retries} retries exhausted. Last error: {last_err}")


def _gemini_review(client, model_name, prompt, max_retries, retry_delay) -> str:
    import google.generativeai as genai
    model = genai.GenerativeModel(
        model_name=model_name,
        system_instruction=SYSTEM_PROMPT_REVIEW,
        generation_config=genai.GenerationConfig(temperature=0.3),
    )
    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            resp = model.generate_content(prompt)
            return resp.text.strip()
        except Exception as e:
            last_err = e
            err = str(e).lower()
            if "429" in str(e) or "quota" in err or "rate" in err:
                wait = retry_delay * attempt
                st.warning(f"⏳ Gemini rate limit (attempt {attempt}/{max_retries}). Waiting {wait}s…")
                time.sleep(wait)
            else:
                raise
    raise Exception(f"Gemini: all {max_retries} retries exhausted. Last error: {last_err}")


# ── Unified dispatcher ───────────────────────────────────────────────────────
def call_ai_redraft(provider, client, model, prompt, max_retries, retry_delay) -> RCMBatchResponse:
    if provider == "Groq":
        return _groq_redraft(client, model, prompt, max_retries, retry_delay)
    else:
        return _gemini_redraft(client, model, prompt, max_retries, retry_delay)


def call_ai_review(provider, client, model, prompt, max_retries, retry_delay) -> str:
    if provider == "Groq":
        return _groq_review(client, model, prompt, max_retries, retry_delay)
    else:
        return _gemini_review(client, model, prompt, max_retries, retry_delay)


# ---------------------------------------------------------------------------
# Initialize client
# ---------------------------------------------------------------------------
def get_client(provider, api_key):
    if provider == "Groq":
        from groq import Groq
        return Groq(api_key=api_key)
    else:
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        return genai  # module itself acts as the configured client


def test_api_key(provider, api_key) -> bool:
    try:
        if provider == "Groq":
            from groq import Groq
            c = Groq(api_key=api_key)
            c.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[{"role": "user", "content": "hi"}],
                max_tokens=5,
            )
        else:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            m = genai.GenerativeModel("gemini-2.0-flash-lite")
            m.generate_content("hi")
        return True
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Export helpers — PDF and Word (unchanged from original)
# ---------------------------------------------------------------------------
def _parse_md_lines(md_text: str):
    lines  = md_text.splitlines()
    parsed = []
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith("### "):
            parsed.append(("h3", s[4:].strip()))
        elif s.startswith("## "):
            parsed.append(("h2", s[3:].strip()))
        elif s.startswith("# "):
            parsed.append(("h1", s[2:].strip()))
        elif s.startswith("---") and all(c == "-" for c in s):
            parsed.append(("hr", ""))
        elif s.startswith("|"):
            cells = [c.strip() for c in s.strip("|").split("|")]
            if not all(set(c.replace("-","").replace(":","").replace(" ","")) == set() for c in cells):
                parsed.append(("table_row", cells))
        elif s.startswith("- ") or s.startswith("* "):
            parsed.append(("bullet", s[2:].strip()))
        elif s.startswith("**") and s.endswith("**") and len(s) > 4:
            parsed.append(("bold_line", s.strip("*")))
        else:
            parsed.append(("normal", s))
    return parsed


def generate_pdf(md_text: str, process: str, industry: str, erp: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    import io as _io

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2.5*cm, bottomMargin=2*cm)
    base = getSampleStyleSheet()
    sty = {
        "h1":     ParagraphStyle("h1",     parent=base["Heading1"], fontSize=16, textColor=colors.HexColor("#1B3A6B"), spaceAfter=10, spaceBefore=14),
        "h2":     ParagraphStyle("h2",     parent=base["Heading2"], fontSize=13, textColor=colors.HexColor("#1B3A6B"), spaceAfter=8,  spaceBefore=12),
        "h3":     ParagraphStyle("h3",     parent=base["Heading3"], fontSize=11, textColor=colors.HexColor("#2E5FA3"), spaceAfter=6,  spaceBefore=10),
        "normal": ParagraphStyle("normal", parent=base["Normal"],   fontSize=9,  leading=13, spaceAfter=4),
        "bullet": ParagraphStyle("bullet", parent=base["Normal"],   fontSize=9,  leading=13, leftIndent=14, bulletIndent=4, spaceAfter=3),
        "bold":   ParagraphStyle("bold",   parent=base["Normal"],   fontSize=9,  leading=13, spaceAfter=4),
        "meta":   ParagraphStyle("meta",   parent=base["Normal"],   fontSize=9,  textColor=colors.HexColor("#555555"), spaceAfter=6),
    }
    story = []
    story.append(Paragraph("RCM Design Effectiveness Review", sty["h1"]))
    story.append(Paragraph(f"<b>Process:</b> {process} &nbsp;|&nbsp; <b>Industry:</b> {industry} &nbsp;|&nbsp; <b>ERP:</b> {erp}", sty["meta"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1B3A6B"), spaceAfter=12))

    parsed = _parse_md_lines(md_text)
    i = 0
    while i < len(parsed):
        typ, content = parsed[i]
        if typ == "h1":
            story.append(Spacer(1, 6)); story.append(Paragraph(content, sty["h1"]))
        elif typ == "h2":
            story.append(Spacer(1, 4)); story.append(Paragraph(content, sty["h2"]))
        elif typ == "h3":
            story.append(Paragraph(content, sty["h3"]))
        elif typ == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#AAAAAA"), spaceBefore=4, spaceAfter=4))
        elif typ == "bullet":
            story.append(Paragraph(f"• {content}", sty["bullet"]))
        elif typ == "bold_line":
            story.append(Paragraph(f"<b>{content}</b>", sty["bold"]))
        elif typ == "table_row":
            rows = [content]
            j = i + 1
            while j < len(parsed) and parsed[j][0] == "table_row":
                rows.append(parsed[j][1]); j += 1
            i = j - 1
            col_count = max(len(r) for r in rows)
            tdata     = [r + [""] * (col_count - len(r)) for r in rows]
            col_width = (A4[0] - 4*cm) / col_count
            tbl = Table(tdata, colWidths=[col_width]*col_count, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (-1, 0), colors.HexColor("#1B3A6B")),
                ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
                ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE",      (0, 0), (-1, -1), 8),
                ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2F8")]),
                ("GRID",          (0, 0), (-1, -1), 0.4, colors.HexColor("#AAAAAA")),
                ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING",    (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(Spacer(1, 4)); story.append(tbl); story.append(Spacer(1, 6))
        elif typ == "normal":
            story.append(Paragraph(content.replace("**", ""), sty["normal"]))
        i += 1

    doc.build(story)
    return buf.getvalue()


def generate_docx(md_text: str, process: str, industry: str, erp: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io as _io, re as _re

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    h = doc.add_heading("RCM Design Effectiveness Review", level=1)
    if h.runs: h.runs[0].font.color.rgb = RGBColor(0x1B, 0x3A, 0x6B)
    meta = doc.add_paragraph()
    for label, val in [("Process: ", process), ("   |   Industry: ", industry), ("   |   ERP: ", erp)]:
        r = meta.add_run(label); r.bold = label.strip().endswith(":")
        meta.add_run(val)

    def add_hr(doc):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "1B3A6B")
        pBdr.append(bottom); pPr.append(pBdr)

    def inline_bold(para, text):
        for part in _re.split(r"(\*\*.*?\*\*)", text):
            if part.startswith("**") and part.endswith("**"):
                para.add_run(part[2:-2]).bold = True
            else:
                para.add_run(part)

    add_hr(doc)
    parsed = _parse_md_lines(md_text)
    i = 0
    while i < len(parsed):
        typ, content = parsed[i]
        if typ in ("h1","h2","h3"):
            lvl = int(typ[1])
            h2 = doc.add_heading(content, level=lvl)
            col = RGBColor(0x1B, 0x3A, 0x6B) if lvl <= 2 else RGBColor(0x2E, 0x5F, 0xA3)
            if h2.runs: h2.runs[0].font.color.rgb = col
        elif typ == "hr":
            add_hr(doc)
        elif typ == "bullet":
            p = doc.add_paragraph(style="List Bullet"); inline_bold(p, content)
        elif typ == "bold_line":
            p = doc.add_paragraph(); p.add_run(content).bold = True
        elif typ == "table_row":
            rows = [content]
            j = i + 1
            while j < len(parsed) and parsed[j][0] == "table_row":
                rows.append(parsed[j][1]); j += 1
            i = j - 1
            col_count = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=col_count)
            tbl.style = "Table Grid"
            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < col_count:
                        cell = tbl.rows[r_idx].cells[c_idx]
                        cell.text = cell_text
                        if r_idx == 0:
                            if cell.paragraphs[0].runs:
                                cell.paragraphs[0].runs[0].bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            tc_pr = cell._tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
                            shd.set(qn("w:fill"), "1B3A6B"); tc_pr.append(shd)
            doc.add_paragraph()
        elif typ == "normal":
            p = doc.add_paragraph(); inline_bold(p, content)
        i += 1

    buf = _io.BytesIO(); doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Main UI
# ---------------------------------------------------------------------------
if not api_key:
    provider_name = "Groq" if provider == "Groq" else "Gemini"
    links = {
        "Groq":            "https://console.groq.com/keys",
        "Gemini (Google)": "https://aistudio.google.com/app/apikey",
    }
    st.info(f"👈 Enter your **{provider_name} API Key** in the sidebar to begin.\n\nGet a free key at 👉 {links[provider]}")
    st.stop()

with st.spinner(f"🔑 Validating {provider} API key…"):
    if not test_api_key(provider, api_key):
        st.error(f"❌ Invalid or expired {provider} API key. Please check and try again.")
        st.stop()

st.success(f"✅ {provider} API key valid! Using model: `{selected_model}`")
st.info(
    f"📋 **Process:** {process_label}  |  "
    f"🏭 **Industry:** {industry_label}  |  "
    f"💻 **ERP:** {erp_label}  |  "
    f"🤖 **Provider:** {provider}  |  "
    f"🛠️ **Mode:** {tool_mode.split('  ')[1]}"
)

client = get_client(provider, api_key)

# ===========================================================================
# MODE A — Redraft & Classify Controls
# ===========================================================================
if "Redraft" in tool_mode:
    st.markdown("## ✏️ Redraft & Classify Controls")
    st.markdown(
        "Upload your RCM Excel file. Each control will be rewritten in **Big 4 audit style** "
        "and classified for **Key/Non-Key** and **Anti-Fraud** using ICFR/SOX criteria."
    )
    uploaded_file = st.file_uploader("📁 Upload RCM Excel File (.xlsx)", type=["xlsx"])

    if uploaded_file:
        df = pd.read_excel(uploaded_file)
        st.write("### 👀 Preview of Uploaded Data")
        st.dataframe(df.head(10))

        st.write("### 🔧 Column Configuration")
        col1, col2 = st.columns(2)
        with col1:
            control_col = st.selectbox("Select the Control Activity Column:", df.columns)
        with col2:
            risk_col = st.selectbox("Select the Risk Column (optional):", ["None"] + list(df.columns))

        if risk_col != "None":
            st.success(f"✅ Risk column `{risk_col}` will be used to improve AI assessments.")

        if st.button("🚀 Process RCM", type="primary"):
            all_results: dict[int, dict] = {}
            rows         = list(df.iterrows())
            batches      = [rows[i: i + batch_size] for i in range(0, len(rows), batch_size)]
            progress_bar = st.progress(0)
            status_text  = st.empty()
            error_log    = []
            has_risk_col = risk_col != "None"

            for b_idx, batch in enumerate(batches):
                status_text.text(f"Processing batch {b_idx + 1} of {len(batches)}…")
                controls_input = []
                for idx, row in batch:
                    item = {"original_id": int(idx), "description": str(row[control_col])}
                    if has_risk_col:
                        item["existing_risk"] = str(row[risk_col])
                    controls_input.append(item)

                prompt = build_redraft_prompt(
                    controls_input, has_risk_col,
                    process_label, industry_label, erp_label
                )

                try:
                    result = call_ai_redraft(
                        provider, client, selected_model, prompt,
                        int(max_retries), retry_delay
                    )
                    for item in result.items:
                        all_results[item.original_id] = {
                            "Redrafted Control Activity": item.redrafted_control,
                            "Anti - Fraud (Y/N)":         item.anti_fraud,
                            "Risk Rating (L,M,H)":        item.risk_rating,
                            "Control Type (Key/Non-Key)": item.control_type,
                        }
                except Exception as e:
                    error_log.append(f"Batch {b_idx + 1}: {e}")
                    st.warning(f"⚠️ Batch {b_idx + 1} failed and was skipped.")

                progress_bar.progress((b_idx + 1) / len(batches))
                if b_idx < len(batches) - 1:
                    time.sleep(2)

            status_text.text("✅ Processing complete!")
            if error_log:
                with st.expander("⚠️ Errors"):
                    for err in error_log: st.write(err)

            if all_results:
                res_df   = pd.DataFrame.from_dict(all_results, orient="index")
                df_clean = df.drop(columns=[c for c in res_df.columns if c in df.columns])
                final_df = pd.concat([df_clean, res_df], axis=1)

                st.success(f"🎉 Done! {len(all_results)} of {len(df)} rows processed.")

                key_count   = sum(1 for v in all_results.values() if v["Control Type (Key/Non-Key)"] == "Key")
                fraud_count = sum(1 for v in all_results.values() if v["Anti - Fraud (Y/N)"] == "Y")
                high        = sum(1 for v in all_results.values() if v["Risk Rating (L,M,H)"].lower() == "high")
                med         = sum(1 for v in all_results.values() if v["Risk Rating (L,M,H)"].lower() == "medium")
                low         = sum(1 for v in all_results.values() if v["Risk Rating (L,M,H)"].lower() == "low")

                c1, c2, c3, c4 = st.columns(4)
                c1.metric("Total Processed",    len(all_results))
                c2.metric("Key Controls",        key_count)
                c3.metric("Non-Key Controls",    len(all_results) - key_count)
                c4.metric("Anti-Fraud Controls", fraud_count)

                st.write("#### 📊 Risk Rating Breakdown")
                r1, r2, r3 = st.columns(3)
                r1.metric("🔴 High", high); r2.metric("🟡 Medium", med); r3.metric("🟢 Low", low)

                st.dataframe(final_df)

                buffer = io.BytesIO()
                with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
                    final_df.to_excel(writer, index=False, sheet_name="Redrafted RCM")
                st.download_button(
                    label="📥 Download Updated RCM Excel",
                    data=buffer.getvalue(),
                    file_name="Redrafted_RCM.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            else:
                st.error("❌ No rows were processed. Try reducing batch size or increasing retry delay.")


# ===========================================================================
# MODE B — Full RCM Design Effectiveness Review
# ===========================================================================
else:
    st.markdown("## 🔍 Full RCM Design Effectiveness Review")
    st.markdown(
        "Upload your RCM Excel file. The AI will perform a comprehensive **14-part "
        "Design Effectiveness Review** benchmarked against COSO 2013, IFC (Companies Act 2013), "
        "ICAI Guidance Note, SOX Best Practices, and Anti-Fraud Framework."
    )
    if not selected_parts:
        st.warning("⚠️ Please select at least one report part in the sidebar.")

    uploaded_file = st.file_uploader("📁 Upload RCM Excel File (.xlsx)", type=["xlsx"])

    if uploaded_file:
        df = pd.read_excel(uploaded_file)
        st.write("### 👀 Preview of Uploaded RCM")
        st.dataframe(df.head(10))

        # ── Step 1: Map key columns explicitly ─────────────────────────────
        st.write("### 🔧 Step 1: Map Your RCM Columns")
        st.caption(
            "Tell the AI exactly which column contains your Risks and which contains your Controls. "
            "This prevents the AI from treating existing controls as missing."
        )

        all_cols     = list(df.columns)
        none_option  = ["— Not in my RCM —"]

        mc1, mc2 = st.columns(2)
        with mc1:
            risk_col_review = st.selectbox(
                "📌 Risk / Risk Description column *",
                none_option + all_cols,
                help="The column that describes what can go wrong (the risk statement)."
            )
        with mc2:
            control_col_review = st.selectbox(
                "📌 Control Activity / Control Description column *",
                none_option + all_cols,
                help="The column that describes the existing control. The AI will review these — NOT treat them as missing."
            )

        mc3, mc4, mc5 = st.columns(3)
        with mc3:
            owner_col_review = st.selectbox(
                "Control Owner column (optional)",
                none_option + all_cols,
            )
        with mc4:
            freq_col_review = st.selectbox(
                "Frequency column (optional)",
                none_option + all_cols,
            )
        with mc5:
            ctrl_type_col_review = st.selectbox(
                "Control Type column (optional)",
                none_option + all_cols,
            )

        # ── Step 2: Additional columns ──────────────────────────────────────
        st.write("### 🔧 Step 2: Additional Columns to Include")
        mapped_cols = [
            c for c in [risk_col_review, control_col_review, owner_col_review,
                        freq_col_review, ctrl_type_col_review]
            if c != "— Not in my RCM —"
        ]
        remaining_cols = [c for c in all_cols if c not in mapped_cols]
        extra_cols = st.multiselect(
            "Any other columns to send to the AI (e.g. Risk ID, Control ID, Assertions):",
            remaining_cols,
            default=[],
        )

        # Validate mandatory fields
        missing_mandatory = []
        if risk_col_review    == "— Not in my RCM —": missing_mandatory.append("Risk column")
        if control_col_review == "— Not in my RCM —": missing_mandatory.append("Control Activity column")

        if missing_mandatory:
            st.warning(f"⚠️ Please map the mandatory columns: {', '.join(missing_mandatory)}")
        else:
            parts_display = "All 14 Parts" if "All Parts (Full Report)" in selected_parts else ", ".join(selected_parts)
            st.info(f"📑 **Report scope:** {parts_display}  |  **Controls:** {len(df)}  |  **Provider:** {provider}")

            if st.button("🔍 Run Design Effectiveness Review", type="primary"):

                # ── Build a clearly-labelled RCM text block ─────────────────
                # Each row is rendered as explicit key: value pairs so the AI
                # can never confuse which field is the risk vs the control.
                def build_labelled_rcm(df, risk_col, control_col, owner_col,
                                       freq_col, ctrl_type_col, extra_cols):
                    lines = []
                    lines.append(
                        "IMPORTANT: The 'Control Activity' field below contains the EXISTING "
                        "control already implemented by the organisation. Do NOT list these as "
                        "missing controls. Evaluate them for design quality, completeness, and "
                        "effectiveness. Only flag controls as MISSING if they address a risk that "
                        "has NO corresponding Control Activity in this RCM.\n"
                    )
                    lines.append("=" * 70)
                    for idx, row in df.iterrows():
                        lines.append(f"\n--- RCM Row #{idx + 1} ---")
                        lines.append(f"Risk Description    : {row[risk_col]}")
                        lines.append(f"Control Activity    : {row[control_col]}")
                        if owner_col and owner_col != "— Not in my RCM —":
                            lines.append(f"Control Owner       : {row[owner_col]}")
                        if freq_col and freq_col != "— Not in my RCM —":
                            lines.append(f"Frequency           : {row[freq_col]}")
                        if ctrl_type_col and ctrl_type_col != "— Not in my RCM —":
                            lines.append(f"Control Type        : {row[ctrl_type_col]}")
                        for ec in extra_cols:
                            lines.append(f"{ec:<20}: {row[ec]}")
                    lines.append("\n" + "=" * 70)
                    return "\n".join(lines)

                rcm_text = build_labelled_rcm(
                    df,
                    risk_col        = risk_col_review,
                    control_col     = control_col_review,
                    owner_col       = owner_col_review,
                    freq_col        = freq_col_review,
                    ctrl_type_col   = ctrl_type_col_review,
                    extra_cols      = extra_cols,
                )

                review_prompt = build_design_review_prompt(
                    process=process_label, industry=industry_label,
                    erp=erp_label, parts=selected_parts, rcm_text=rcm_text,
                )

                with st.spinner("🔍 AI is reviewing your RCM — this may take 1–2 minutes…"):
                    try:
                        review_output = call_ai_review(
                            provider, client, selected_model, review_prompt,
                            int(max_retries), retry_delay,
                        )
                        st.success("✅ Design Effectiveness Review complete!")
                        st.markdown("---")
                        st.markdown("## 📋 RCM Design Effectiveness Review Report")
                        st.markdown(f"**Process:** {process_label}  |  **Industry:** {industry_label}  |  **ERP:** {erp_label}  |  **Provider:** {provider}")
                        st.markdown("---")
                        st.markdown(review_output)

                        st.markdown("### 📥 Download Report")
                        dl1, dl2, dl3 = st.columns(3)
                        fname_base = f"RCM_Review_{process_label.replace(' ','_')}"

                        with dl1:
                            st.download_button("📄 Markdown (.md)", data=review_output.encode("utf-8"),
                                               file_name=f"{fname_base}.md", mime="text/markdown",
                                               use_container_width=True)
                        with dl2:
                            try:
                                docx_bytes = generate_docx(review_output, process_label, industry_label, erp_label)
                                st.download_button("📝 Word (.docx)", data=docx_bytes,
                                                   file_name=f"{fname_base}.docx",
                                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                                   use_container_width=True)
                            except Exception as ex:
                                st.warning(f"Word export failed: {ex}")
                        with dl3:
                            try:
                                pdf_bytes = generate_pdf(review_output, process_label, industry_label, erp_label)
                                st.download_button("📕 PDF (.pdf)", data=pdf_bytes,
                                                   file_name=f"{fname_base}.pdf", mime="application/pdf",
                                                   use_container_width=True)
                            except Exception as ex:
                                st.warning(f"PDF export failed: {ex}")

                    except Exception as e:
                        st.error(f"❌ Review failed: {e}")
